"""Verify the normalized customer-file corpus, provenance, and asset links."""

from __future__ import annotations

import argparse
import csv
import hashlib
import json
import re
import zipfile
from collections import Counter
from pathlib import Path
from typing import Any

from docx import Document


OS_METADATA = {".DS_Store", "Thumbs.db"}
EXPECTED_INVENTORY_FILES = 108
EXPECTED_DOCUMENTS = 20
EXPECTED_OS_METADATA = 4
EXPECTED_EMBEDDED_MEDIA = 21
EXPECTED_RENDERED_PAGES = 9


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        while chunk := handle.read(1024 * 1024):
            digest.update(chunk)
    return digest.hexdigest()


def read_csv(path: Path) -> list[dict[str, str]]:
    with path.open(encoding="utf-8-sig", newline="") as handle:
        return list(csv.DictReader(handle))


def local_markdown_links(path: Path) -> list[Path]:
    content = path.read_text(encoding="utf-8")
    links: list[Path] = []
    for target in re.findall(r"\[[^\]]*\]\(([^)]+)\)", content):
        if target.startswith(("http://", "https://", "mailto:", "tel:", "#")):
            continue
        clean = target.split("#", 1)[0]
        if clean:
            links.append((path.parent / clean).resolve())
    return links


def normalized_text(value: str) -> str:
    return re.sub(r"\s+", " ", value).strip()


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("customer_root", type=Path)
    parser.add_argument("repo_root", type=Path)
    args = parser.parse_args()

    customer_root = args.customer_root.resolve()
    repo = args.repo_root.resolve()
    output = repo / "docs/source-material/customer-files"
    errors: list[str] = []

    inventory_data: dict[str, Any] = json.loads((output / "inventory.json").read_text(encoding="utf-8"))
    inventory = inventory_data.get("files", [])
    if inventory_data.get("file_count") != EXPECTED_INVENTORY_FILES or len(inventory) != EXPECTED_INVENTORY_FILES:
        errors.append(f"inventory must contain {EXPECTED_INVENTORY_FILES} files")
    if sum(int(item["bytes"]) for item in inventory) != int(inventory_data.get("total_bytes", -1)):
        errors.append("inventory byte total does not reconcile")

    ids = [item["id"] for item in inventory]
    paths = [item["source_path"] for item in inventory]
    hashes = [item["sha256"] for item in inventory]
    if len(ids) != len(set(ids)):
        errors.append("inventory contains duplicate stable IDs")
    if len(paths) != len(set(paths)):
        errors.append("inventory contains duplicate source paths")
    if len(hashes) != len(set(hashes)):
        errors.append("inventory contains duplicate hashes")

    os_rows = [item for item in inventory if Path(item["source_path"]).name in OS_METADATA]
    substantive = [item for item in inventory if Path(item["source_path"]).name not in OS_METADATA]
    if len(os_rows) != EXPECTED_OS_METADATA:
        errors.append(f"expected {EXPECTED_OS_METADATA} OS metadata records, found {len(os_rows)}")
    if any(item["cleanup_status"] != "removed-os-metadata" for item in os_rows):
        errors.append("OS metadata rows do not have removed-os-metadata status")
    if any(item["cleanup_status"] != "preserve-original" for item in substantive):
        errors.append("a substantive source is not marked preserve-original")

    current_files = [path for path in customer_root.rglob("*") if path.is_file()]
    forbidden = [path for path in current_files if path.name in OS_METADATA]
    if forbidden:
        errors.extend(f"OS metadata still present: {path.relative_to(repo).as_posix()}" for path in forbidden)

    inventory_by_path = {item["source_path"]: item for item in substantive}
    current_substantive = [path for path in current_files if path.name not in OS_METADATA]
    if len(current_substantive) != len(substantive):
        errors.append(f"substantive file count mismatch: inventory {len(substantive)}, current {len(current_substantive)}")
    for path in current_substantive:
        relative = path.relative_to(customer_root).as_posix()
        item = inventory_by_path.get(relative)
        if item is None:
            errors.append(f"current substantive file missing from inventory: {relative}")
            continue
        if path.stat().st_size != int(item["bytes"]):
            errors.append(f"source size changed: {relative}")
            continue
        if sha256(path) != item["sha256"]:
            errors.append(f"source hash changed: {relative}")

    prior = read_csv(repo / "docs/source-material/asset-manifest.csv")
    prior_hashes = {row["sha256"]: row for row in prior if row.get("sha256")}
    for item in substantive:
        previous = prior_hashes.get(item["sha256"])
        if previous is None:
            errors.append(f"substantive source missing from prior manifest: {item['source_path']}")
            continue
        if item["prior_status"] != previous["status"]:
            errors.append(f"prior status mismatch: {item['source_path']}")
        if item["curated_destination"] != previous["destination"]:
            errors.append(f"curated destination mismatch: {item['source_path']}")
        destination = item["curated_destination"]
        if destination and not (repo / destination).is_file():
            errors.append(f"missing curated destination: {destination}")

    document_rows = read_csv(output / "document-index.csv")
    if len(document_rows) != EXPECTED_DOCUMENTS:
        errors.append(f"expected {EXPECTED_DOCUMENTS} normalized documents, found {len(document_rows)}")
    if Counter(row["extraction_status"] for row in document_rows) != Counter({"ok": EXPECTED_DOCUMENTS}):
        errors.append("not every normalized document has extraction_status=ok")
    inventory_by_id = {item["id"]: item for item in inventory}
    visual_content = json.loads((output / "visual-content-transcriptions.json").read_text(encoding="utf-8"))
    embedded_transcriptions = visual_content.get("embedded_images", {})
    pdf_visual_summaries = visual_content.get("pdf_visual_summaries", {})
    if len(embedded_transcriptions) != EXPECTED_EMBEDDED_MEDIA:
        errors.append(
            f"expected {EXPECTED_EMBEDDED_MEDIA} embedded-image transcriptions, found {len(embedded_transcriptions)}"
        )
    exported_embedded = 0
    rendered_pages = 0
    for row in document_rows:
        item = inventory_by_id.get(row["id"])
        if item is None:
            errors.append(f"document references unknown inventory ID: {row['id']}")
            continue
        normalized = repo / row["normalized_path"]
        if not normalized.is_file() or normalized.stat().st_size < 300:
            errors.append(f"missing or empty normalized document: {row['normalized_path']}")
            continue
        content = normalized.read_text(encoding="utf-8")
        if item["sha256"] not in content or item["source_path"] not in content:
            errors.append(f"normalized provenance mismatch: {row['normalized_path']}")
        if any(marker in content for marker in ("Ã", "Â", "â€")):
            errors.append(f"mojibake marker in normalized document: {row['normalized_path']}")

        source = customer_root / row["source_path"]
        media_dir = output / "media" / row["id"]
        if source.suffix.lower() == ".docx":
            document = Document(source)
            for paragraph in document.paragraphs:
                paragraph_text = paragraph.text.strip()
                if paragraph_text and paragraph_text not in content:
                    errors.append(f"DOCX paragraph missing from normalized text: {row['source_path']}: {paragraph_text[:80]}")
            for table in document.tables:
                for table_row in table.rows:
                    for cell in table_row.cells:
                        cell_text = cell.text.strip()
                        if cell_text and cell_text not in content:
                            errors.append(f"DOCX table text missing from normalized text: {row['source_path']}: {cell_text[:80]}")
            with zipfile.ZipFile(source) as archive:
                package_media = [
                    name
                    for name in archive.namelist()
                    if name.startswith("word/media/") and not name.endswith("/")
                ]
            exported = sorted(media_dir.glob("embedded-*")) if media_dir.exists() else []
            if len(exported) != len(package_media) or len(exported) != int(row.get("embedded_media") or 0):
                errors.append(f"DOCX embedded-media count mismatch: {row['source_path']}")
            exported_embedded += len(exported)
            for asset in exported:
                digest = sha256(asset)
                if digest not in embedded_transcriptions:
                    errors.append(f"embedded image has no visual transcription: {asset.relative_to(repo).as_posix()}")
                else:
                    visible_lines = [
                        normalized_text(line)
                        for line in embedded_transcriptions[digest].get("visible_text", "").splitlines()
                        if normalized_text(line)
                    ]
                    if any(line not in normalized_text(content) for line in visible_lines):
                        errors.append(f"embedded transcription missing from Markdown: {row['normalized_path']}")
        elif source.suffix.lower() == ".pdf":
            pages = int(row.get("pages") or 0)
            renders = sorted(media_dir.glob("page-*.png")) if media_dir.exists() else []
            if len(renders) != pages or len(renders) != int(row.get("rendered_pages") or 0):
                errors.append(f"PDF rendered-page count mismatch: {row['source_path']}")
            rendered_pages += len(renders)
            if int(row.get("pages_with_text") or 0) == 0:
                summary = pdf_visual_summaries.get(item["sha256"])
                if not summary:
                    errors.append(f"image/vector-only PDF lacks a visual summary: {row['source_path']}")
                else:
                    for visible_text in summary.get("visible_text", []):
                        if visible_text not in content:
                            errors.append(f"PDF visual summary missing from Markdown: {row['normalized_path']}")

    if exported_embedded != EXPECTED_EMBEDDED_MEDIA:
        errors.append(f"expected {EXPECTED_EMBEDDED_MEDIA} exported DOCX media files, found {exported_embedded}")
    if rendered_pages != EXPECTED_RENDERED_PAGES:
        errors.append(f"expected {EXPECTED_RENDERED_PAGES} rendered PDF pages, found {rendered_pages}")

    cross_reference = read_csv(output / "asset-cross-reference.csv")
    cross_ids = Counter(row["id"] for row in cross_reference)
    if set(cross_ids) != set(ids):
        errors.append("cross-reference IDs do not cover the inventory exactly")
    if any(count < 1 for count in cross_ids.values()):
        errors.append("an inventory item has no cross-reference row")
    for row in cross_reference:
        item = inventory_by_id.get(row["id"])
        if item is None:
            continue
        if row["source_path"] != item["source_path"] or row["sha256"] != item["sha256"]:
            errors.append(f"cross-reference provenance mismatch: {row['id']}")
        for field in ("curated_destination", "public_destination"):
            destination = row[field]
            if destination and not (repo / destination).is_file():
                errors.append(f"missing {field}: {destination}")
        for usage in filter(None, (part.strip() for part in row["code_usage"].split(";"))):
            if not (repo / usage).is_file():
                errors.append(f"missing code-usage path: {usage}")

    derivative_rows = read_csv(output / "public-derivative-map.csv")
    for row in derivative_rows:
        if row["source_path"] not in inventory_by_path:
            errors.append(f"derivative map source is missing: {row['source_path']}")
        if not (repo / row["public_path"]).is_file():
            errors.append(f"derivative map public file is missing: {row['public_path']}")

    structures = json.loads((output / "page-structure-reference.json").read_text(encoding="utf-8"))
    if not structures.get("entries"):
        errors.append("page-structure reference has no entries")
    for entry in structures.get("entries", []):
        for source_id in entry.get("source_document_ids", []):
            if source_id not in inventory_by_id:
                errors.append(f"page structure references unknown source ID: {source_id}")
        for normalized_path in entry.get("normalized_documents", []):
            if not (output / normalized_path).is_file():
                errors.append(f"page structure normalized document is missing: {normalized_path}")
        for evidence in entry.get("current_evidence", []):
            if not (repo / evidence).exists():
                errors.append(f"page structure evidence is missing: {evidence}")

    for markdown in output.rglob("*.md"):
        for link in local_markdown_links(markdown):
            if not link.exists():
                errors.append(f"broken local Markdown link in {markdown.relative_to(repo).as_posix()}: {link}")

    if errors:
        print("FAILED")
        for error in errors:
            print(f"- {error}")
        raise SystemExit(1)

    print(f"PASS: {len(inventory)} inventory records reconcile to {len(substantive)} preserved sources + {len(os_rows)} removed OS metadata files")
    print(f"PASS: all {len(substantive)} substantive customer files match recorded sizes and SHA-256 hashes")
    print(f"PASS: all {len(document_rows)} DOCX/PDF sources have non-empty normalized Markdown")
    print(f"PASS: {exported_embedded} embedded DOCX images have hash-linked visual transcriptions")
    print(f"PASS: all {rendered_pages} PDF pages are rendered; image/vector-only PDFs have visual summaries")
    print(f"PASS: {len(cross_reference)} source/curated/public cross-reference rows resolve")
    print(f"PASS: {len(structures['entries'])} page-structure entries resolve to source and implementation evidence")
    print("PASS: no OS metadata remains and local Markdown links resolve")


if __name__ == "__main__":
    main()
