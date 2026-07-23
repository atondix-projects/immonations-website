"""Build a machine-readable corpus and provenance index for customer-files/."""

from __future__ import annotations

import argparse
import csv
import hashlib
import json
import re
import shutil
import subprocess
import sys
import unicodedata
import zipfile
from collections import defaultdict
from datetime import date
from pathlib import Path
from typing import Any

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from lxml import etree
from PIL import Image, ImageOps
from pypdf import PdfReader


OS_METADATA = {".DS_Store", "Thumbs.db"}
DOCUMENT_EXTENSIONS = {".docx", ".pdf"}
IMAGE_EXTENSIONS = {".jpg", ".jpeg", ".png", ".webp", ".heic", ".svg"}
VIDEO_EXTENSIONS = {".mp4", ".mov", ".m4v", ".webm"}
SEARCH_EXTENSIONS = {".ts", ".tsx", ".js", ".jsx", ".md", ".mdx", ".json", ".html", ".txt"}
WORD_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
OFFICE_REL_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
PACKAGE_REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        while chunk := handle.read(1024 * 1024):
            digest.update(chunk)
    return digest.hexdigest()


def slugify(value: str) -> str:
    value = value.translate(str.maketrans({"ä": "ae", "ö": "oe", "ü": "ue", "Ä": "ae", "Ö": "oe", "Ü": "ue", "ß": "ss"}))
    value = unicodedata.normalize("NFKD", value).encode("ascii", "ignore").decode("ascii")
    return re.sub(r"[^a-z0-9]+", "-", value.lower()).strip("-") or "file"


def markdown_cell(value: str) -> str:
    return re.sub(r"\s+", " ", value).strip().replace("|", "\\|")


def paragraph_markdown(paragraph: Paragraph) -> str:
    text = paragraph.text.strip()
    if not text:
        return ""

    style = paragraph.style.name if paragraph.style else ""
    heading = re.match(r"(?:Heading|Überschrift)\s*(\d+)", style, re.IGNORECASE)
    if heading:
        level = min(int(heading.group(1)) + 1, 6)
        return f"{'#' * level} {text}"

    properties = paragraph._p.pPr
    if properties is not None and properties.numPr is not None:
        return f"- {text}"

    return text


def table_markdown(table: Table) -> str:
    rows = [[markdown_cell(cell.text) for cell in row.cells] for row in table.rows]
    if not rows:
        return ""
    width = max(len(row) for row in rows)
    rows = [row + [""] * (width - len(row)) for row in rows]
    header = rows[0]
    body = rows[1:]
    lines = [f"| {' | '.join(header)} |", f"| {' | '.join(['---'] * width)} |"]
    lines.extend(f"| {' | '.join(row)} |" for row in body)
    return "\n".join(lines)


def serializable_properties(document: Document) -> dict[str, str]:
    properties = document.core_properties
    result: dict[str, str] = {}
    for name in ("title", "subject", "author", "keywords", "comments", "category", "created", "modified", "last_modified_by"):
        value = getattr(properties, name, None)
        if value not in (None, ""):
            result[name] = value.isoformat() if hasattr(value, "isoformat") else str(value)
    return result


def package_part_text(archive: zipfile.ZipFile, prefix: str) -> list[str]:
    blocks: list[str] = []
    namespace = {"w": WORD_NS}
    for name in sorted(archive.namelist()):
        if not name.startswith(prefix) or not name.endswith(".xml"):
            continue
        root = etree.fromstring(archive.read(name))
        text = " ".join(part.strip() for part in root.xpath("//w:t/text()", namespaces=namespace) if part.strip())
        if text and text not in blocks:
            blocks.append(text)
    return blocks


def docx_hyperlinks(archive: zipfile.ZipFile) -> list[dict[str, str]]:
    relationship_path = "word/_rels/document.xml.rels"
    if relationship_path not in archive.namelist():
        return []
    rel_root = etree.fromstring(archive.read(relationship_path))
    targets = {
        rel.get("Id", ""): rel.get("Target", "")
        for rel in rel_root.findall(f"{{{PACKAGE_REL_NS}}}Relationship")
        if (rel.get("Type") or "").endswith("/hyperlink")
    }
    document_root = etree.fromstring(archive.read("word/document.xml"))
    namespace = {"w": WORD_NS, "r": OFFICE_REL_NS}
    links: list[dict[str, str]] = []
    for node in document_root.xpath("//w:hyperlink", namespaces=namespace):
        relationship_id = node.get(f"{{{OFFICE_REL_NS}}}id", "")
        display = "".join(node.xpath(".//w:t/text()", namespaces=namespace)).strip()
        target = targets.get(relationship_id, "")
        anchor = node.get(f"{{{WORD_NS}}}anchor", "")
        if target or anchor:
            links.append({"display": display or target or anchor, "target": target or f"#{anchor}"})
    return links


def docx_drawing_count(archive: zipfile.ZipFile) -> int:
    root = etree.fromstring(archive.read("word/document.xml"))
    namespace = {"w": WORD_NS}
    return len(root.xpath("//w:drawing | //w:pict", namespaces=namespace))


def export_docx_media(
    archive: zipfile.ZipFile,
    media_dir: Path,
    markdown_prefix: str,
    transcriptions: dict[str, dict[str, Any]],
) -> list[dict[str, Any]]:
    names = sorted(name for name in archive.namelist() if name.startswith("word/media/") and not name.endswith("/"))
    assets: list[dict[str, Any]] = []
    if names:
        media_dir.mkdir(parents=True, exist_ok=True)
    for index, package_name in enumerate(names, start=1):
        extension = Path(package_name).suffix.lower() or ".bin"
        destination = media_dir / f"embedded-{index:02d}{extension}"
        destination.write_bytes(archive.read(package_name))
        details = image_metadata(destination) if extension in IMAGE_EXTENSIONS else {}
        digest = sha256(destination)
        assets.append(
            {
                "package_name": package_name,
                "exported_path": f"{markdown_prefix}/{destination.name}",
                "bytes": destination.stat().st_size,
                "sha256": digest,
                "transcription": transcriptions.get(digest),
                **details,
            }
        )
    return assets


def extract_docx(
    path: Path,
    media_dir: Path,
    markdown_prefix: str,
    transcriptions: dict[str, dict[str, Any]],
) -> tuple[str, dict[str, Any]]:
    document = Document(path)
    blocks: list[str] = []
    paragraphs = 0
    tables = 0
    for item in document.iter_inner_content():
        if isinstance(item, Paragraph):
            rendered = paragraph_markdown(item)
            if rendered:
                paragraphs += 1
                blocks.append(rendered)
        elif isinstance(item, Table):
            rendered = table_markdown(item)
            if rendered:
                tables += 1
                blocks.append(rendered)
    with zipfile.ZipFile(path) as archive:
        headers = package_part_text(archive, "word/header")
        footers = package_part_text(archive, "word/footer")
        hyperlinks = docx_hyperlinks(archive)
        drawings = docx_drawing_count(archive)
        media = export_docx_media(archive, media_dir, markdown_prefix, transcriptions)

    if headers:
        blocks.extend(["## Headers", *(f"- {text}" for text in headers)])
    if footers:
        blocks.extend(["## Footers", *(f"- {text}" for text in footers)])
    if hyperlinks:
        blocks.extend(["## Hyperlinks", *(f"- [{link['display']}]({link['target']})" for link in hyperlinks)])
    if media:
        blocks.append("## Embedded media")
        for index, asset in enumerate(media, start=1):
            dimensions = ""
            if asset.get("width") and asset.get("height"):
                dimensions = f"; {asset['width']}x{asset['height']} px"
            blocks.extend(
                [
                    f"### Embedded image {index}",
                    f"- Package source: `{asset['package_name']}`",
                    f"- SHA-256: `{asset['sha256']}`",
                    f"- Size: {asset['bytes']} bytes{dimensions}",
                    f"![Embedded source image {index}]({asset['exported_path']})",
                ]
            )
            transcription = asset.get("transcription")
            if transcription:
                blocks.append("#### Visible-text transcription")
                for label, key in (
                    ("Content type", "content_type"),
                    ("Reviewer", "reviewer"),
                    ("Profile context", "profile_context"),
                    ("Rating", "rating"),
                    ("Relative date", "relative_date"),
                    ("Visible badge", "visible_badge"),
                ):
                    if transcription.get(key) not in (None, ""):
                        suffix = "/5" if key == "rating" else ""
                        blocks.append(f"- {label}: {transcription[key]}{suffix}")
                visible_text = str(transcription.get("visible_text", "")).strip()
                if visible_text:
                    blocks.append("\n".join(">" if not line else f"> {line}" for line in visible_text.splitlines()))
                for note_key in ("transcription_note", "interface_note"):
                    if transcription.get(note_key):
                        blocks.append(f"- Note: {transcription[note_key]}")

    return "\n\n".join(blocks), {
        "paragraphs": paragraphs,
        "tables": tables,
        "headers": len(headers),
        "footers": len(footers),
        "hyperlinks": len(hyperlinks),
        "drawings": drawings,
        "embedded_media": len(media),
        "transcribed_embedded_media": sum(1 for asset in media if asset.get("transcription")),
        "document_properties": serializable_properties(document),
    }


def extract_pdf(path: Path, media_dir: Path, markdown_prefix: str, pdftoppm: str) -> tuple[str, dict[str, Any]]:
    reader = PdfReader(path)
    pages: list[str] = []
    nonempty = 0
    media_dir.mkdir(parents=True, exist_ok=True)
    for index, page in enumerate(reader.pages, start=1):
        text = (page.extract_text() or "").strip()
        if text:
            nonempty += 1
        render_path = media_dir / f"page-{index:02d}.png"
        command = [
            pdftoppm,
            "-f",
            str(index),
            "-l",
            str(index),
            "-singlefile",
            "-scale-to",
            "1600",
            "-png",
            str(path),
            str(render_path.with_suffix("")),
        ]
        subprocess.run(command, check=True, capture_output=True)
        pages.append(
            "\n\n".join(
                [
                    f"## Page {index}",
                    f"![Rendered source page {index}]({markdown_prefix}/{render_path.name})",
                    text or "[No extractable text on this page; the rendered page above preserves the visual/vector content.]",
                ]
            )
        )
    metadata = {str(key).lstrip("/"): str(value) for key, value in (reader.metadata or {}).items() if value not in (None, "")}
    return "\n\n".join(pages), {
        "pages": len(reader.pages),
        "pages_with_text": nonempty,
        "rendered_pages": len(reader.pages),
        "pdf_metadata": metadata,
    }


def image_metadata(path: Path) -> dict[str, Any]:
    if path.suffix.lower() == ".svg":
        return {"format": "SVG"}
    try:
        with Image.open(path) as image:
            image = ImageOps.exif_transpose(image)
            return {"format": image.format, "width": image.width, "height": image.height}
    except Exception as error:
        return {"metadata_error": f"{type(error).__name__}: {error}"}


def video_metadata(path: Path, ffprobe: str) -> dict[str, Any]:
    command = [
        ffprobe,
        "-v",
        "error",
        "-show_entries",
        "format=duration:stream=codec_type,codec_name,width,height",
        "-of",
        "json",
        str(path),
    ]
    try:
        result = subprocess.run(command, check=True, capture_output=True, text=True, encoding="utf-8")
        data = json.loads(result.stdout)
        duration = data.get("format", {}).get("duration")
        streams = data.get("streams", [])
        video = next((stream for stream in streams if stream.get("codec_type") == "video"), {})
        audio = next((stream for stream in streams if stream.get("codec_type") == "audio"), {})
        return {
            "duration_seconds": round(float(duration), 3) if duration else None,
            "video_codec": video.get("codec_name"),
            "width": video.get("width"),
            "height": video.get("height"),
            "audio_codec": audio.get("codec_name"),
        }
    except Exception as error:
        return {"metadata_error": f"{type(error).__name__}: {error}"}


def load_prior_manifest(path: Path) -> dict[str, dict[str, str]]:
    with path.open(encoding="utf-8-sig", newline="") as handle:
        rows = list(csv.DictReader(handle))
    return {row["sha256"]: row for row in rows if row.get("sha256")}


def load_derivative_map(path: Path) -> dict[str, list[dict[str, str]]]:
    if not path.exists():
        return {}
    with path.open(encoding="utf-8-sig", newline="") as handle:
        rows = list(csv.DictReader(handle))
    result: dict[str, list[dict[str, str]]] = defaultdict(list)
    for row in rows:
        result[row["source_path"]].append(row)
    return result


def source_usage(repo: Path, public_path: str) -> list[str]:
    needle = "/" + public_path.replace("\\", "/").removeprefix("public/")
    matches: list[str] = []
    for root_name in ("src", "messages", "content", "docs"):
        root = repo / root_name
        if not root.exists():
            continue
        for path in root.rglob("*"):
            if not path.is_file() or path.suffix.lower() not in SEARCH_EXTENSIONS:
                continue
            if "source-material" in path.parts:
                continue
            try:
                content = path.read_text(encoding="utf-8")
            except UnicodeDecodeError:
                continue
            if needle in content:
                matches.append(path.relative_to(repo).as_posix())
    return sorted(matches)


def write_csv(path: Path, rows: list[dict[str, Any]], fields: list[str]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=fields, extrasaction="ignore")
        writer.writeheader()
        writer.writerows(rows)


def resolve_pdftoppm(value: str) -> str:
    if value != "pdftoppm":
        return value
    executable = "pdftoppm.exe" if sys.platform == "win32" else "pdftoppm"
    bundled = Path(sys.executable).resolve().parent.parent / "native" / "poppler" / "Library" / "bin" / executable
    return str(bundled) if bundled.is_file() else value


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("customer_root", type=Path)
    parser.add_argument("repo_root", type=Path)
    parser.add_argument("--ffprobe", default="ffprobe")
    parser.add_argument("--pdftoppm", default="pdftoppm")
    args = parser.parse_args()
    pdftoppm = resolve_pdftoppm(args.pdftoppm)

    customer_root = args.customer_root.resolve()
    repo = args.repo_root.resolve()
    output = repo / "docs/source-material/customer-files"
    corpus = output / "documents"
    media_root = output / "media"
    corpus.mkdir(parents=True, exist_ok=True)
    if media_root.exists():
        if media_root.parent != output or media_root.name != "media":
            raise RuntimeError(f"Refusing unexpected media cleanup path: {media_root}")
        shutil.rmtree(media_root)
    media_root.mkdir(parents=True)

    prior = load_prior_manifest(repo / "docs/source-material/asset-manifest.csv")
    derivative_map = load_derivative_map(output / "public-derivative-map.csv")
    visual_content_path = output / "visual-content-transcriptions.json"
    visual_content = json.loads(visual_content_path.read_text(encoding="utf-8"))
    embedded_transcriptions = visual_content.get("embedded_images", {})
    pdf_visual_summaries = visual_content.get("pdf_visual_summaries", {})
    existing_inventory_path = output / "inventory.json"
    removed_os_records: list[dict[str, Any]] = []
    if existing_inventory_path.exists():
        existing_inventory = json.loads(existing_inventory_path.read_text(encoding="utf-8")).get("files", [])
        removed_os_records = [
            item
            for item in existing_inventory
            if Path(item["source_path"]).name in OS_METADATA
            and not (customer_root / Path(item["source_path"])).exists()
        ]
    existing_public_hashes: dict[str, list[str]] = defaultdict(list)
    for path in (repo / "public").rglob("*"):
        if path.is_file():
            existing_public_hashes[sha256(path)].append(path.relative_to(repo).as_posix())

    inventory: list[dict[str, Any]] = []
    cross_reference: list[dict[str, Any]] = []
    document_index: list[dict[str, Any]] = []

    files = sorted((path for path in customer_root.rglob("*") if path.is_file()), key=lambda path: path.as_posix().casefold())
    for path in files:
        relative = path.relative_to(customer_root).as_posix()
        digest = sha256(path)
        extension = path.suffix.lower()
        stable_id = f"cf-{digest[:12]}"
        is_os_metadata = path.name in OS_METADATA
        previous = prior.get(digest, {})
        metadata: dict[str, Any] = {}

        if extension in IMAGE_EXTENSIONS:
            metadata = image_metadata(path)
        elif extension in VIDEO_EXTENSIONS:
            metadata = video_metadata(path, args.ffprobe)
        elif extension == ".pdf":
            metadata = {"pages": len(PdfReader(path).pages)}
        elif extension == ".docx":
            document = Document(path)
            metadata = {"paragraphs": len(document.paragraphs), "tables": len(document.tables)}

        record = {
            "id": stable_id,
            "source_path": relative,
            "extension": extension,
            "bytes": path.stat().st_size,
            "sha256": digest,
            "category": relative.split("/", 1)[0],
            "cleanup_status": "remove-os-metadata" if is_os_metadata else "preserve-original",
            "metadata": metadata,
            "prior_status": previous.get("status", "unmatched"),
            "curated_destination": previous.get("destination", ""),
        }
        inventory.append(record)

        normalized_path = ""
        extraction_status = "not-a-document"
        if extension in DOCUMENT_EXTENSIONS:
            try:
                document_media = media_root / stable_id
                markdown_media_prefix = f"../media/{stable_id}"
                if extension == ".docx":
                    text, extraction_metadata = extract_docx(
                        path,
                        document_media,
                        markdown_media_prefix,
                        embedded_transcriptions,
                    )
                else:
                    text, extraction_metadata = extract_pdf(
                        path,
                        document_media,
                        markdown_media_prefix,
                        pdftoppm,
                    )
                    visual_summary = pdf_visual_summaries.get(digest)
                    if visual_summary:
                        visible_text = "; ".join(visual_summary.get("visible_text", []))
                        text = "\n\n".join(
                            [
                                text,
                                "## Visual-only content transcription",
                                f"- Content type: {visual_summary.get('content_type', 'visual content')}",
                                f"- Visible text: {visible_text}",
                                f"- Visual description: {visual_summary.get('visual_description', '')}",
                            ]
                        )
                        extraction_metadata["visual_summary"] = visual_summary
                metadata.update(extraction_metadata)
                extraction_status = "ok" if text.strip() else "empty"
            except Exception as error:
                text = f"[Extraction error: {type(error).__name__}: {error}]"
                extraction_metadata = {}
                extraction_status = "error"
            filename = f"{slugify(path.stem)}-{digest[:8]}.md"
            normalized = corpus / filename
            normalized_path = normalized.relative_to(repo).as_posix()
            frontmatter = {
                "id": stable_id,
                "source_path": relative,
                "source_sha256": digest,
                "source_bytes": path.stat().st_size,
                "source_type": extension.lstrip("."),
                "extraction_status": extraction_status,
                **extraction_metadata,
            }
            metadata_lines = ["---", *(f"{key}: {json.dumps(value, ensure_ascii=False)}" for key, value in frontmatter.items()), "---"]
            normalized.write_text(
                "\n".join([*metadata_lines, "", f"# {path.stem}", "", "> Supplied customer source. Treat claims and copy as unapproved until verified.", "", text, ""]),
                encoding="utf-8",
            )
            document_index.append(
                {
                    "id": stable_id,
                    "source_path": relative,
                    "normalized_path": normalized_path,
                    "extraction_status": extraction_status,
                    **{
                        key: json.dumps(value, ensure_ascii=False, sort_keys=True)
                        if isinstance(value, (dict, list))
                        else value
                        for key, value in extraction_metadata.items()
                    },
                }
            )

        exact_public = existing_public_hashes.get(digest, [])
        manual_derivatives = derivative_map.get(relative, [])
        if exact_public:
            for public_path in exact_public:
                cross_reference.append(
                    {
                        "id": stable_id,
                        "source_path": relative,
                        "sha256": digest,
                        "curated_destination": previous.get("destination", ""),
                        "public_destination": public_path,
                        "match_basis": "exact-sha256",
                        "confidence": "exact",
                        "code_usage": "; ".join(source_usage(repo, public_path)),
                        "notes": "Byte-identical public file.",
                    }
                )
        elif manual_derivatives:
            for derivative in manual_derivatives:
                public_path = derivative["public_path"]
                cross_reference.append(
                    {
                        "id": stable_id,
                        "source_path": relative,
                        "sha256": digest,
                        "curated_destination": previous.get("destination", ""),
                        "public_destination": public_path,
                        "match_basis": derivative["match_basis"],
                        "confidence": derivative["confidence"],
                        "code_usage": "; ".join(source_usage(repo, public_path)),
                        "notes": derivative["notes"],
                    }
                )
        else:
            cross_reference.append(
                {
                    "id": stable_id,
                    "source_path": relative,
                    "sha256": digest,
                    "curated_destination": previous.get("destination", ""),
                    "public_destination": "",
                    "match_basis": "no-public-derivative-identified",
                    "confidence": "n/a",
                    "code_usage": "",
                    "notes": "Curated source may still be reusable; no verified public derivative is recorded.",
                }
            )

    for item in removed_os_records:
        item["cleanup_status"] = "removed-os-metadata"
        inventory.append(item)
        cross_reference.append(
            {
                "id": item["id"],
                "source_path": item["source_path"],
                "sha256": item["sha256"],
                "curated_destination": "",
                "public_destination": "",
                "match_basis": "removed-os-metadata",
                "confidence": "exact",
                "code_usage": "",
                "notes": "Operating-system metadata removed during source cleanup; inventory record retained for provenance.",
            }
        )

    inventory.sort(key=lambda item: item["source_path"].casefold())
    cross_reference.sort(key=lambda item: (item["source_path"].casefold(), item["public_destination"]))

    inventory_path = output / "inventory.json"
    inventory_path.write_text(
        json.dumps(
            {
                "schema_version": 1,
                "generated_on": date.today().isoformat(),
                "source_root": "customer-files",
                "file_count": len(inventory),
                "total_bytes": sum(item["bytes"] for item in inventory),
                "files": inventory,
            },
            ensure_ascii=False,
            indent=2,
        )
        + "\n",
        encoding="utf-8",
    )
    write_csv(
        output / "inventory.csv",
        [
            {
                **item,
                "metadata": json.dumps(item["metadata"], ensure_ascii=False, sort_keys=True),
            }
            for item in inventory
        ],
        ["id", "source_path", "extension", "bytes", "sha256", "category", "cleanup_status", "metadata", "prior_status", "curated_destination"],
    )
    write_csv(
        output / "document-index.csv",
        document_index,
        [
            "id",
            "source_path",
            "normalized_path",
            "extraction_status",
            "paragraphs",
            "tables",
            "headers",
            "footers",
            "hyperlinks",
            "drawings",
            "embedded_media",
            "transcribed_embedded_media",
            "pages",
            "pages_with_text",
            "rendered_pages",
            "document_properties",
            "pdf_metadata",
            "visual_summary",
        ],
    )
    write_csv(
        output / "asset-cross-reference.csv",
        cross_reference,
        ["id", "source_path", "sha256", "curated_destination", "public_destination", "match_basis", "confidence", "code_usage", "notes"],
    )
    print(f"Indexed {len(inventory)} files ({sum(item['bytes'] for item in inventory)} bytes)")
    print(f"Normalized {len(document_index)} documents to {corpus}")
    print(f"Wrote {len(cross_reference)} cross-reference rows")


if __name__ == "__main__":
    main()
