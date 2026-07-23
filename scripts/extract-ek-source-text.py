"""Extract reviewable text from the EK source archive without altering originals."""

from __future__ import annotations

import argparse
import re
import zipfile
from pathlib import Path

from docx import Document
from pypdf import PdfReader


def slugify(value: str) -> str:
    replacements = str.maketrans(
        {
            "ä": "ae",
            "ö": "oe",
            "ü": "ue",
            "Ä": "ae",
            "Ö": "oe",
            "Ü": "ue",
            "ß": "ss",
        }
    )
    value = value.translate(replacements).lower()
    value = re.sub(r"[^a-z0-9]+", "-", value).strip("-")
    return value or "document"


def extract_docx(path: Path) -> str:
    document = Document(path)
    blocks: list[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            blocks.append(text)

    for table_index, table in enumerate(document.tables, start=1):
        blocks.append(f"\n[TABLE {table_index}]")
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " / ") for cell in row.cells]
            blocks.append(" | ".join(cells))

    return "\n\n".join(blocks)


def extract_pdf(path: Path) -> str:
    reader = PdfReader(path)
    pages: list[str] = []
    for index, page in enumerate(reader.pages, start=1):
        pages.append(f"[PAGE {index}]\n{page.extract_text() or '[NO EXTRACTABLE TEXT]'}")
    return "\n\n".join(pages)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("archive", type=Path)
    parser.add_argument("output", type=Path)
    args = parser.parse_args()

    args.output.mkdir(parents=True, exist_ok=True)
    extracted_dir = args.output / "originals"
    extracted_dir.mkdir(exist_ok=True)

    with zipfile.ZipFile(args.archive) as archive:
        entries = [
            entry
            for entry in archive.infolist()
            if not entry.is_dir() and Path(entry.filename).suffix.lower() in {".docx", ".pdf"}
        ]

        for entry in entries:
            relative_parts = Path(entry.filename).parts[1:]
            stem = slugify("--".join(Path(part).stem for part in relative_parts))
            source_path = extracted_dir / f"{stem}{Path(entry.filename).suffix.lower()}"
            source_path.write_bytes(archive.read(entry))

            try:
                text = extract_docx(source_path) if source_path.suffix == ".docx" else extract_pdf(source_path)
                status = "ok" if text.strip() else "empty"
            except Exception as error:  # Preserve the source and make failure reviewable.
                text = f"[EXTRACTION ERROR] {type(error).__name__}: {error}"
                status = "error"

            review_path = args.output / f"{stem}.md"
            review_path.write_text(
                "\n".join(
                    [
                        f"# {Path(entry.filename).name}",
                        "",
                        f"- Source: `{entry.filename}`",
                        f"- Extraction status: `{status}`",
                        f"- Source bytes: `{entry.file_size}`",
                        "",
                        "## Extracted text",
                        "",
                        text,
                        "",
                    ]
                ),
                encoding="utf-8",
            )

    print(f"Extracted {len(entries)} documents to {args.output}")


if __name__ == "__main__":
    main()
